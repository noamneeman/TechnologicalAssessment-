from abc import ABC, abstractmethod
from io import BytesIO

import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from matplotlib.figure import Figure
from tqdm import tqdm
import os
from docx.oxml.ns import qn
from column_constants import SIGMAS
from docxtpl import DocxTemplate
from docx.oxml import OxmlElement

ADD_IN_END_OF_SENTENCE: str = "."


def text_to_rgba(s, *, dpi, **kwargs):
    # To convert a text string to an image, we can:
    # - draw it on an empty and transparent figure;
    # - save the figure to a temporary buffer using ``bbox_inches="tight",
    #   pad_inches=0`` which will pick the correct area to save;
    # - load the buffer using ``plt.imread``.
    #
    # (If desired, one can also directly save the image to the filesystem.)
    fig = Figure(facecolor="none")
    fig.text(0, 0, s, **kwargs)
    with BytesIO() as buf:
        fig.savefig(buf, dpi=dpi, format="png", bbox_inches="tight",
                    pad_inches=0)
        buf.seek(0)
        rgba = plt.imread(buf)
    return rgba

def fix_rtl_symbols(sentence):
    """Fix punctuation in RTL text using BiDi marks."""
    if not isinstance(sentence, str):
        return sentence

    # Add RLM around punctuation
    rtl_marks = "\u200F"
    sentence = sentence.replace(",", f"{rtl_marks},{rtl_marks}")
    sentence = sentence.replace(".", f"{rtl_marks}.{rtl_marks}")
    sentence = sentence.replace("\"", f"{rtl_marks}\"{rtl_marks}")
    sentence = sentence.replace("\\", f"{rtl_marks}\\{rtl_marks}")
    sentence = sentence.replace("-", f"{rtl_marks}-{rtl_marks}")
    sentence = sentence.replace(":", f"{rtl_marks}:{rtl_marks}")
    sentence = sentence.replace(";", f"{rtl_marks};{rtl_marks}")
    sentence = sentence.replace("(", f"{rtl_marks}({rtl_marks}")
    sentence = sentence.replace(")", f"{rtl_marks}){rtl_marks}")

    # Wrap entire sentence with RLE and PDF
    sentence = f"\u202B{sentence}\u202C"
    return sentence


def add_in_the_beginning(sentence, add_on):
    if isinstance(sentence, str) and sentence[0] != add_on:
        sentence = add_on + sentence
    return sentence


def ensure_ends_with(sentence, add_on):
    if isinstance(sentence, str):
        if sentence[-1] == " ":
            sentence = sentence[:-1]
        if sentence[-1] != add_on:
            sentence += add_on
    sentence = fix_rtl_symbols(sentence)
    return sentence


class Docx_helper(ABC):
    def __init__(self, file_format_path, word_output_dir):
        self.file_format_path = file_format_path
        self.word_output_dir = word_output_dir

    @abstractmethod
    def is_values(self, category: str) -> bool:
        pass

    @abstractmethod
    def get_columns_names(self):
        pass

    def dehash(self, s: str):
        return "".join([chr(ord(c) - i % 5) for i, c in enumerate(s)])

    def my_hash(self, s: str):
        # increase the utf of each char by 1
        return "".join([chr(ord(c) + i % 5) for i, c in enumerate(s)])

    def sigma_text(self, sigma, category):
        values_small_threshold, values_big_threshold = SIGMAS[category]

        small, big = False, False
        if sigma > values_big_threshold:
            big = True
        elif sigma < values_small_threshold:
            small = True

        if big:
            ret_val = "sigma is top 15% (big)"
        elif small:
            ret_val = "sigma is lowest 15% (small)"
        else:
            ret_val = "sigma value is average"
        return ret_val

    def create_histogram(self, all_avgs, avg_total, avg_personal, std_personal, category,
                         old_average=-1, N=-1):
        fig = plt.figure()
        ax = plt.gca()

        is_values = self.is_values(category)

        if is_values:
            bins = np.arange(-0.125, 3.375, 0.25)
            xticks = np.arange(0, 3.25, 0.25)
        else:
            bins = np.arange(-0.25, 6.75, 0.5)
            xticks = np.arange(0, 6.5, 0.5)

        plt.xticks(xticks, fontsize=16)
        plt.hist(x=all_avgs, bins=bins, rwidth=0.9)
        plt.yticks(fontsize=16)

        # plot the average value of the specific person
        plt.axvline(avg_personal, color='red')

        # plot the average value of person form last year
        if old_average != -1:
            plt.axvline(old_average, color='green', linestyle="--")
            plt.text(0.01, 0.7, s="Red line - new result\nDashed line - last semester", fontsize=12, color='black',
                     transform=ax.transAxes)

        # text of number of comments for this category (N)
        if N != -1:
            plt.text(0.01, 0.6, s=f"N (none zero) ={N}", fontsize=16, color='blue', transform=ax.transAxes)

        # plot the std of the specific person
        # the name column is the index, so we need the i'th column
        plt.hlines(y=sum(ax.get_ylim()) / 2, xmin=avg_personal - std_personal,
                   xmax=avg_personal + std_personal, color='red')
        # plt.text(0.01, 0.93, transform=ax.transAxes,
        #          s=r'$\sigma$' + f'={std_personal}\n{self.sigma_text(std_personal, is_values)}',
        #          fontsize=16, color='red')
        plt.text(0.01, 0.93, transform=ax.transAxes, s=r'$\sigma$' + f"={std_personal}", fontsize=16, color='red')
        plt.text(0.01, 0.89, s=self.sigma_text(std_personal, category), fontsize=16, color='red',
                 transform=ax.transAxes)

        plt.axvline(avg_total, color='black')
        secondary_ax = ax.secondary_xaxis("top")
        # plotting the value of the axvline on the histogram
        if abs(avg_personal - avg_total) < 0.2:
            diff = (0.2 - abs(avg_personal - avg_total)) / 2
            if avg_total > avg_personal:
                secondary_ax.set_xticks(ticks=[avg_personal - diff, avg_total + diff],
                                        labels=[f"{round(avg_personal, 2)}",
                                                f"{round(avg_total, 2)}"], rotation=60)
            else:
                secondary_ax.set_xticks(ticks=[avg_total - diff, avg_personal + diff],
                                        labels=[f"{round(avg_total, 2)}",
                                                f"{round(avg_personal, 2)}"], rotation=60)
        else:
            secondary_ax.set_xticks(ticks=[avg_personal, avg_total],
                                    labels=[f"{round(avg_personal, 2)}", f"{round(avg_total, 2)}"],
                                    rotation=60)

        for label in secondary_ax.get_xticklabels():
            label.set_fontsize(16)

        fig.set_size_inches(10, 5)
        plt.close(fig)
        # fig.show()
        # plt.show()
        return fig

    def insert_classifications(self, classification_df, format_file_name):
        conserve_names = [("Interpersonal Skills", "יכולות בין-אישיות"), \
                          ("Intrapersonal Skills", "יכולות תוך-אישיות"), \
                          ("Professionalism", "מקצועיות"), \
                          ("Conduct", "התנהלות"), \
                          ("Leadership", "מנהיגות"), \
                          ("Other", "אחר")]

        improve_names = [("Interpersonal Skills2", "יכולות בין-אישיות"), \
                         ("Intrapersonal Skills2", "יכולות תוך-אישיות"), \
                         ("Professionalism2", "מקצועיות"), \
                         ("Conduct2", "התנהלות"), \
                         ("Leadership2", "מנהיגות"), \
                         ("Other2", "אחר")]

        rtl_marks = "\u200F"
        doc = DocxTemplate(template_file=format_file_name)
        context = {'conserve_classifications': [], 'improve_classifications': []}

        for column, word_name in conserve_names:
            list_of_sentences = []
            if column not in classification_df:
                continue

            for i in range(len(classification_df[column])):
                if classification_df[column][i] in ["True", True, "TRUE", "true"]:
                    current_sentence = classification_df['Original_conserve'][i]
                    current_sentence = ensure_ends_with(current_sentence, ADD_IN_END_OF_SENTENCE)
                    list_of_sentences.append({'name': current_sentence})

            if len(list_of_sentences) > 0:
                class_dict = {'name': f"\u202B {word_name} {rtl_marks}({rtl_marks}{len(list_of_sentences)}{rtl_marks}){rtl_marks} " + ":\u202C"}
                class_dict["bullets"] = list_of_sentences

                context["conserve_classifications"].append(class_dict)

        for column, word_name in improve_names:
            list_of_sentences = []
            if column not in classification_df:
                continue

            for i in range(len(classification_df[column])):
                if classification_df[column][i] in ["True", True, "TRUE", "true"]:
                    current_sentence = classification_df['Original_improve'][i]
                    current_sentence = ensure_ends_with(current_sentence, ADD_IN_END_OF_SENTENCE)
                    list_of_sentences.append({'name': current_sentence})

            if len(list_of_sentences) > 0:
                class_dict = {'name': f"\u202B {word_name} {rtl_marks}({rtl_marks}{len(list_of_sentences)}{rtl_marks}){rtl_marks} " + ":\u202C"}
                class_dict["bullets"] = list_of_sentences

                context["improve_classifications"].append(class_dict)

        doc.render(context=context)
        doc.save(format_file_name)

    def create_word_file(self, hists, classification_df, person_name, n=None, names_to_hashes=False):
        if names_to_hashes:
            with open("names_to_hashes.txt", "a", encoding="utf-8") as f:
                f.write(f"{person_name} => {self.my_hash(person_name)}\n")
            title_to_save = f"{self.my_hash(person_name)} (N={n})".replace('"', '').replace("'", '') + ".docx"
        else:
            title_to_save = f"{person_name} (N={n})".replace('"', '').replace("'", '') + ".docx"
        path_to_save = os.path.join(self.word_output_dir, title_to_save)

        doc = Document(self.file_format_path)

        title = doc.paragraphs[0]
        if names_to_hashes:
            title.text += self.my_hash(person_name)
        else:
            title.text += person_name
        if n is not None:
            title.text += f" (N={n})"
        title.runs[0] = "David"
        title.runs[0].underline = True
        title.runs[0].style.font.size = Pt(12)

        # add the histograms to the table
        for i in range(len(hists)):
            # save the hist as png, so we can load to word as a picture
            cur_hist = hists[i]
            TMP_FILE_PATH = "tmp.png"
            cur_hist.savefig(TMP_FILE_PATH)

            # load png
            cell = doc.tables[0].cell(i, 1)
            cell.add_paragraph().add_run().add_picture(TMP_FILE_PATH, height=Inches(1.9))
            cell.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # remove png as it is no longer needed
            os.remove(TMP_FILE_PATH)

        # for paragraph in doc.paragraphs:
        #     self.set_paragraph_rtl(paragraph)

        # avoiding corrupting the word file
        # the id gets mixed with some of the things of the template file
        # for further reading - https://github.com/python-openxml/python-docx/issues/455
        # or need version better than 0.8.7
        try:
            docPrs = doc._part._element.findall('.//' + qn('wp:docPr'))
            for docPr in docPrs:
                docPr.set('id', str(int(docPr.get('id')) + 100000))
        except:
            pass

        doc.save(path_to_save)

        if classification_df is not None:
            self.insert_classifications(classification_df, path_to_save)

    def set_paragraph_rtl(self, paragraph):
        # Set paragraph alignment to right
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Set paragraph direction to RTL
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    def create_text_figure(self, avg_personal):
        val_to_sentence = {1: "נמוך ביחס לממוצע ",
         2: "מתחת לממוצע",
         3: "מעט מתחת לממוצע",
         4: "מעט מעל הממוצע",
         5: "מעל הממוצע",
         6: "גבוה ביחס לממוצע"}
        fig = plt.figure()
        rgba2 = text_to_rgba(r""+f"{val_to_sentence[round(avg_personal)][::-1]}", color="black", fontsize=15, dpi=200)
        fig.figimage(rgba2, 200, 300)
        return fig

    def run_word_creation(self,
                          combined_df: pd.DataFrame,
                          stats_df: pd.DataFrame,
                          name_to_classification: dict = None,
                          old_stats_df=None,
                          verbose: bool = True,
                          start_cadet: str = None,
                          names_to_hashes: bool = False):

        reached_start_cadet = start_cadet is None
        # create word file for every person
        for df in tqdm(combined_df.groupby("name"), disable=not verbose):
            person_name = df[0]
            if not reached_start_cadet:
                if person_name == start_cadet:
                    reached_start_cadet = True
                else:
                    continue

            df = df[1]

            # TODO - voodoo code to get only numerical columns
            numerical_columns = df.columns.drop("name")[:-3]  # drop the conserve, improve and good talpion columns
            stats_per_person = stats_df[stats_df["name"] == person_name]
            no_hist_numerical_columns = df.columns.drop("name")[-3:-2]

            hists = []
            for category in numerical_columns:
                # calculate old value
                old_average = -1
                if old_stats_df is not None:
                    # does old_stats_df contain the person?
                    if len(old_stats_df[old_stats_df["name"] == person_name]) == 0:
                        print(f"Person {person_name} not found in old stats!\n"
                              f"probably someone changed their name in the raw data excel file\n")

                    old_res = \
                        old_stats_df[(old_stats_df["name"] == person_name) & (old_stats_df["category"] == category)][
                            "mean"]
                    if len(old_res) > 0:
                        old_average = old_res.array[0]

                combined_col = combined_df[category]
                # remove from col every non numeric value row
                combined_col = combined_col[
                    combined_col.apply(lambda x: isinstance(x, (int, np.int64, float, np.float64)))]
                combined_col = combined_col[combined_col.apply(lambda x: x > 0)]  # filter 0 values

                # calculate the number of non 0 comments for this name, for this category
                N = combined_col[combined_df["name"] == person_name].shape[0]

                avg_total = stats_df[stats_df["category"] == category]["mean"].mean() # mean of all the people
                # avg_total = combined_col.mean()

                avg_personal = stats_per_person[stats_per_person["category"] == category]["mean"].values[0]
                std_personal = stats_per_person[stats_per_person["category"] == category]["std"].values[0]
                all_avgs = stats_df[stats_df["category"] == category]["mean"]

                hist = self.create_histogram(all_avgs, avg_total, avg_personal, std_personal, category,
                                             old_average=old_average, N=N)
                hists.append(hist)

            # Add the colums whom we want only the avrage and std to be presented without the histogram
            for category in no_hist_numerical_columns:
                avg_personal = stats_per_person[stats_per_person["category"] == category]["mean"].values[0]
                hist = self.create_text_figure(avg_personal)
                hists.append(hist)


            N = df.shape[0]
            # print(f"Creating word file for {person_name} (N={N})")
            df = name_to_classification[person_name] if name_to_classification is not None else None

            self.create_word_file(hists, df, person_name, n=N, names_to_hashes=names_to_hashes)

        # save the hash text file to the output directory
        with open(self.word_output_dir + "\\names_to_hashes.txt", "w", encoding="utf-8") as f:
            with open("names_to_hashes.txt", "r", encoding="utf-8") as f2:
                f.write(f2.read())

        # open the temporary hash file and delete all of the content
        with open("names_to_hashes.txt", "w", encoding="utf-8") as f:
            f.write("")

