import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import qn
from docx.shared import Inches, Pt

FORMAT_PATH = "../Formats/demographic_report_format.docx"
OUTPUT_PATH = "output/"


class DocxHelper:
    def __init__(self, title: str):
        self.doc = Document(FORMAT_PATH)
        title_template = self.doc.paragraphs[1]
        title_template.text += title
        title_template.runs[0] = "David"
        title_template.runs[0].underline = True
        title_template.runs[0].bold = True
        # set the title size to 16
        title_template.runs[0].font.size = Pt(16)

    def add_histograms(self, hists, width=6.6, height=2.2):
        """
        add the histograms to the tables, each pair of histograms is added to a row in the tables respectively

        :param height: height of the histograms in inches
        :param width: width of the histograms in inches
        :param hists:
        :return:
        """
        for i in range(int(len(hists))):
            # save the hist as png, so we can load to word as a picture
            cur_hist = hists[i]
            TMP_FILE_PATH = "tmp.png"
            with open(TMP_FILE_PATH, 'wb') as f:
                f.write(cur_hist)

            # load png
            row = int(i / 2)
            if i % 2 == 0:
                cell = self.doc.tables[0].cell(row + 1, 1)
            else:
                cell = self.doc.tables[1].cell(row + 1, 1)
            cell.add_paragraph().add_run().add_picture(TMP_FILE_PATH, width=Inches(width), height=Inches(height))
            cell.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # remove png as it is no longer needed
            os.remove(TMP_FILE_PATH)

        # avoiding corrupting the word file
        # the id gets mixed with some of the things of the template file
        # for further reading - https://github.com/python-openxml/python-docx/issues/455
        # or need version better than 0.8.7
        try:
            docPrs = self.doc._part._element.findall('.//' + qn('wp:docPr'))
            for docPr in docPrs:
                docPr.set('id', str(int(docPr.get('id')) + 100000))
        except:
            pass

    def add_histograms_to_table(self, hists, table_num, width=6.6, height=2.2):
        """
        add the histograms to the specified table

        :param table_num: num of the table to add the histograms to
        :param height: height of the histograms in inches
        :param width: width of the histograms in inches
        :param hists:
        :return:
        """
        for i in range(int(len(hists))):
            # save the hist as png, so we can load to word as a picture
            cur_hist = hists[i]
            TMP_FILE_PATH = "tmp.png"
            with open(TMP_FILE_PATH, 'wb') as f:
                f.write(cur_hist)
            try:
                cell = self.doc.tables[table_num].cell(i + 1, 1)
            except:
                print(f"table {table_num} does not have enough rows")
                print(f"failed at i={i}")
                break
            cell.add_paragraph().add_run().add_picture(TMP_FILE_PATH, width=Inches(width), height=Inches(height))
            cell.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # remove png as it is no longer needed
            os.remove(TMP_FILE_PATH)

        # avoiding corrupting the word file
        # the id gets mixed with some of the things of the template file
        # for further reading - https://github.com/python-openxml/python-docx/issues/455
        # or need version better than 0.8.7
        try:
            docPrs = self.doc._part._element.findall('.//' + qn('wp:docPr'))
            for docPr in docPrs:
                docPr.set('id', str(int(docPr.get('id')) + 100000))
        except:
            pass

    def save(self, path: str):
        self.doc.save(path)
